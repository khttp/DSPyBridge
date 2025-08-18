from docx import Document
import dspy
from app.core import config
from sentence_transformers import SentenceTransformer
import re
from typing import List

# Configure LM
lm = dspy.LM(
    model=config.DEFAULT_MODEL,
    temperature=0.7,
    api_key=config.GROQ_API_KEY,
)
dspy.configure(lm=lm)

# 1. IMPROVED DOCUMENT CHUNKING
def read_docx_with_chunks(filepath, chunk_size=500, overlap=100):
    """Read DOCX and split into overlapping chunks"""
    doc = Document(filepath)
    
    # Extract all text including tables
    all_text = []
    
    # Get paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            all_text.append(p.text.strip())
    
    # Get tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                all_text.append(row_text)
    
    full_text = "\n".join(all_text)
    
    # Split into chunks with overlap
    chunks = []
    words = full_text.split()
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
    
    return chunks

# 2. BETTER EMBEDDING MODEL (more specialized)
def get_better_embedder():
    """Use a better embedding model for retrieval"""
    # Options (choose one):
    
    # Option A: Better general model
    model = SentenceTransformer("all-mpnet-base-v2")  # Better than MiniLM
    
    # Option B: For domain-specific (if available)
    # model = SentenceTransformer("msmarco-distilbert-base-v4")  # Good for Q&A
    
    return model

# 3. QUERY ENHANCEMENT
def enhance_query(question: str) -> str:
    """Enhance the user query for better retrieval"""
    # Add context keywords based on the domain
    domain_keywords = {
        'roaming': ['international', 'travel', 'abroad', 'foreign', 'outside'],
        'extend': ['renewal', 'prolonge', 'continue', 'activate'],
        'date': ['until', 'till', 'expire', 'expiry', 'period'],
        'vodafone': ['network', 'service', 'plan', 'package']
    }
    
    enhanced = question.lower()
    
    # Add relevant keywords based on detected intent
    for key, synonyms in domain_keywords.items():
        if key in enhanced:
            enhanced += f" {' '.join(synonyms[:2])}"  # Add 2 most relevant synonyms
    
    return enhanced

# 4. IMPROVED CORPUS PREPARATION
corpus = []
hf_model = get_better_embedder()

# Load documents with chunking
docx_chunks = read_docx_with_chunks("docs\\Roaming_SummerOffers.docx", chunk_size=300, overlap=50)
corpus.extend(docx_chunks)

print(f"Created {len(corpus)} chunks from document")

# 5. BETTER EMBEDDING FUNCTION WITH NORMALIZATION
def improved_embed(texts):
    """Improved embedding with normalization"""
    if isinstance(texts, str):
        texts = [texts]
    
    # Get embeddings
    embeddings = hf_model.encode(texts, convert_to_numpy=True, normalize_embeddings=True)
    return embeddings.tolist()

# 6. CUSTOM RETRIEVER WITH RERANKING
class ImprovedRetriever:
    def __init__(self, embedder, corpus, k=5):
        self.base_retriever = dspy.retrievers.Embeddings(
            embedder=embedder,
            corpus=corpus,
            k=k*2  # Get more candidates for reranking
        )
        self.corpus = corpus
        self.k = k
    
    def __call__(self, query):
        # Enhance query
        enhanced_query = enhance_query(query)
        
        # Get initial results
        initial_results = self.base_retriever(enhanced_query)
        
        # Simple reranking based on exact keyword matches
        scored_results = []
        query_words = set(query.lower().split())
        
        for doc in initial_results:
            doc_words = set(doc.lower().split())
            exact_matches = len(query_words.intersection(doc_words))
            scored_results.append((exact_matches, doc))
        
        # Sort by exact matches and return top k
        scored_results.sort(key=lambda x: x[0], reverse=True)
        return [doc for score, doc in scored_results[:self.k]]

# 7. IMPROVED RETRIEVER SETUP
retriever = ImprovedRetriever(
    embedder=improved_embed,
    corpus=corpus,
    k=3
)

# 8. BETTER RAG SIGNATURE WITH MORE SPECIFIC INSTRUCTIONS
class ImprovedRAGSignature(dspy.Signature):
    """Answer questions about Vodafone telecommunication services using retrieved context."""
    question = dspy.InputField(desc="User question about Vodafone services (roaming, plans, etc.)")
    context = dspy.InputField(desc="Relevant information from Vodafone documentation")
    answer = dspy.OutputField(desc="Helpful and accurate answer based on the context. If extending roaming, provide specific steps and requirements.")

# 9. ENHANCED RAG MODULE WITH BETTER CONTEXT PROCESSING
class ImprovedRAG(dspy.Module):
    def __init__(self):
        super().__init__()
        self.retrieve = retriever
        self.answer = dspy.Predict(ImprovedRAGSignature)
    
    def forward(self, question):
        # Retrieve top-k docs
        docs = self.retrieve(question)
        
        # Better context formatting
        context_parts = []
        for i, doc in enumerate(docs, 1):
            # Clean up the text
            clean_doc = re.sub(r'\s+', ' ', doc.strip())
            context_parts.append(f"[Context {i}]: {clean_doc}")
        
        context = "\n\n".join(context_parts)
        
        # Add instruction for better answers
        enhanced_question = f"{question}\n\nPlease provide a specific and actionable answer based on the Vodafone documentation."
        
        # Predict answer with retrieved context
        return self.answer(context=context, question=enhanced_question)

# 10. ADVANCED: ADD FALLBACK WITH DIFFERENT RETRIEVAL STRATEGIES
class HybridRAG(dspy.Module):
    def __init__(self):
        super().__init__()
        self.semantic_retrieve = retriever
        self.answer = dspy.Predict(ImprovedRAGSignature)
    
    def keyword_search(self, question, top_k=2):
        """Simple keyword-based fallback"""
        question_words = question.lower().split()
        scored_docs = []
        
        for doc in corpus:
            doc_lower = doc.lower()
            score = sum(doc_lower.count(word) for word in question_words)
            if score > 0:
                scored_docs.append((score, doc))
        
        scored_docs.sort(key=lambda x: x[0], reverse=True)
        return [doc for score, doc in scored_docs[:top_k]]
    
    def forward(self, question):
        # Get semantic results
        semantic_docs = self.semantic_retrieve(question)
        
        # Get keyword results as backup
        keyword_docs = self.keyword_search(question)
        
        # Combine and deduplicate
        all_docs = semantic_docs + [doc for doc in keyword_docs if doc not in semantic_docs]
        
        # Format context
        context_parts = []
        for i, doc in enumerate(all_docs[:4], 1):  # Limit to top 4
            clean_doc = re.sub(r'\s+', ' ', doc.strip())
            context_parts.append(f"[Source {i}]: {clean_doc}")
        
        context = "\n\n".join(context_parts)
        
        return self.answer(context=context, question=question)

# Usage options:
# Option 1: Basic improved version
rag_improved = ImprovedRAG()

# Option 2: Hybrid version (recommended)
rag_hybrid = HybridRAG()

# Test both
question = "Hi i want to extend my roaming till 6/6/2025"

print("=== IMPROVED RAG ===")
response1 = rag_improved(question)
print("Answer:", response1.answer)

print("\n=== HYBRID RAG ===")
response2 = rag_hybrid(question)
print("Answer:", response2.answer)
